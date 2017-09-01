# -*- coding: utf-8 -*-
'''

Created on 2017/08/22

@author: yuyang
'''

import os
import urllib
import uuid
import re
import docx_ext

from docx.shared import Pt
from docx.shared import RGBColor

def addAuthor(document, author):
    para = document.add_paragraph()
    run = para.add_run(author)
    font = run.font
    #font.name = 'Microsoft YaHei'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0x43, 0x6E, 0xEE)
    
def addContent(document, content):
    topic_words = []
    if content.__contains__('k.t.qq.com'):
        pattern = re.compile(r'(<a href="http://k.t.qq.com.*?</a>)', re.S)
        topics = re.findall(pattern, content)
        for topic in topics:
            topic_word = topic.split('#')[1]
            topic_words.append(topic_word)
            content = content.replace(topic, '|||')
    content_parts = content.split('|||')
    
    para = document.add_paragraph()
    for index, content_part in enumerate(content_parts):
        run_c = para.add_run(content_part)
        font_c = run_c.font
        font_c.size = Pt(16)
        font_c.color.rgb = RGBColor(0x08, 0x08, 0x08)
        if len(topic_words) >= index + 1 and len(topic_words) != 0:
            run_t = para.add_run('#' + topic_words[index] + '#')
            font_t = run_t.font
            font_t.italic = True
            font_t.size = Pt(16)
            font_t.color.rgb = RGBColor(0x00, 0x00, 0xCD)

def addPicture(document, story):
    filenames = analyzePic(story)
    for filename in filenames:
        try:
            document.add_picture(filename)
        except:
            print '插入图片出错：' + filename

def addTime(document, time):
    para = document.add_paragraph()
    run = para.add_run(time)
    font = run.font
    font.italic = True
    #font.name = 'Microsoft YaHei'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x7A, 0x7A, 0x7A)
    
def addLocation(document, story):
    location_items = analyzeLoc(story)
    if len(location_items) <= 0:
        return
    link_name = location_items[2]
    google_map_url = 'https://maps.google.com/maps?q=' + location_items[0] + ',' + location_items[1]
    print google_map_url
    para = document.add_paragraph()
    run = para.add_run(u'位置：')
    font = run.font
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x7A, 0x7A, 0x7A)
    docx_ext.add_hyperlink(para, google_map_url, link_name, '4169E1', False)

def downloadPic(url):
    try:
        if not os.path.exists('.//pics'):
            os.mkdir('.//pics')
        filename = '.\\pics\\' + str(uuid.uuid4()) + '.jpg'
        urllib.urlretrieve(url, filename)
    except Exception:
        print 'Download picture error: ' + url
    return filename

def analyzePic(story):
    filenames = []
    if story.__contains__('class="picBox"'):
        pattern = re.compile(r'<div class="picBox">\n<a href="(.*?)" data-like', re.S)
        img_url = re.findall(pattern, story)[0]
        print '图片：',
        print img_url
        filename = downloadPic(img_url)
        filenames.append(filename)
    elif story.__contains__('class="tl_imgGroup'):
        pattern = re.compile(r'<div class="tl_imgGroup(.*?)<div class="miniMultiMedia clear"', re.S)
        imgs_str = re.findall(pattern, story)[0]
        pattern_img = re.compile(r'<a href="(.*?)" class="tl_imgGroup', re.S)
        imgs = re.findall(pattern_img, imgs_str)
        for img_url in imgs:
            print '图片：',
            print img_url
            filename = downloadPic(img_url)
            filenames.append(filename)
            
    return filenames

def analyzeLoc(story):
    location_items = []
    if story.__contains__('class="areaInfo"'):
        pattern = re.compile(r'boss="btn_check_tweetNear".*?lat=(.*?)&lng=(.*?)&addr=(.*?)" target', re.S)
        location_items = re.findall(pattern, story)[0]
        print u'位置：' + location_items[2]
        print u'经度：' + location_items[0]
        print u'纬度：' + location_items[1]

    return location_items
        
    