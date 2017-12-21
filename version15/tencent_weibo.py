# -*- coding: utf-8 -*-
'''
change log:

    修改运行环境为Python3.x
    以对应控制台执行时中文LOG乱码导致异常
    =============OLD=============
    修改文件为分割存储
    加入转帖与视频封面插入图片异常处理
    修改备份文件名生成规则
    对应其它各种表情
    对应",<符号的正确显示
    对应带有转发的内容
    对应 &，>符号的正确显示
    对应含有QQ表情的内容
    对应含有URL的内容
    调整图片为统一宽度
    对应含有@好友的内容
    对应分享视频链接的微博
    添加备份时开始与结束页的控制
    对应含有Emoji的内容
    对应Python的全名规则
    对应含有话题的内容
    对应位置信息(谷歌地图)
    优化生成Word文件格式
    下载图片时文件夹不存在则创建
    对应一条微博多张图片时的图片下载
    对应含有图片的微博
    备份微博内容（仅文字部分）  

Created on 2017/08/10

@author: yuyang
'''
import time
import urllib
from urllib import request
import re
from docx import Document
import tencent_util
#from version15 import  tencent_util

class tencent_weibo:
    '''
    Tencent weibo object.
    '''
    START_PAGE_INDEX = 1
    END_PAGE_INDEX = 3
    SAVE_FILE_PAGE = 20


    def __init__(self, id):
        '''
        Constructor
        '''
        self.id = id
        self.page_index = 1
        self.have_next_page = True
        self.next_page_url = None
        self.user_agent = 'Mozilla/4.0 (compatible; MSIE 5.5; Windows NT)'
        self.headers = {'User-Agent' : self.user_agent}
        self.document = Document()
        
    def get_next_page(self):
        if self.page_index == 1:
            url = 'http://t.qq.com/' + self.id + '?mode=0&lang=zh_CN'
        else:
            url = self.next_page_url + '&lang=zh_CN'
        try:
            req = request.Request(url, headers = self.headers)
            #request = urllib.Request(url, headers = self.headers)
            res = request.urlopen(req)
            page = res.read().decode('utf-8')
        except urllib.URLError as e:
            if hasattr(e, 'reason'):
                print('腾讯微博第 %d 页连接失败： %s' % (self.page_index, e.reason))
                page = None
        
        if page:
            for match in re.finditer(r'<a href="\?mode=0&.*?"', page):
                pass
            self.next_page_url = match.group().replace('<a href="', 'http://t.qq.com/' + self.id)[:-1]
            
            if not page.__contains__(u'下一页'):
                self.have_next_page = False
        
        return page
 
 
    def get_stories(self, page):
        pattern = re.compile(r'<ul id="talkList"(.*?)</ul>', re.S)
        stories_str = re.findall(pattern, page)
        stories = stories_str[0].split('<li')
        return stories
            
    def get_items(self, stories):
        
        pattern = re.compile(r'<div class="msgBox".*?<div class="userName".*?title="(.*?)" gender=' +
                             '.*?<div class="msgCnt">(.*?)</div>' +
                             '.*?<div class="pubInfo.*?from="\\d*">(.*?)</a>', re.S)
        
        for story in stories:
            story, quotation = tencent_util.depart_quotation(story)
            
            items = re.findall(pattern, story)
            if not len(items):
                continue
            
            item = items[0]
            print('作者：', item[0])
            self.document.add_heading('', 0)
            tencent_util.add_author(self.document, item[0])
            print('内容：', item[1])
            content_valid = ''.join(c for c in item[1] if ord(c) >= 32)# delete control chars
            tencent_util.add_content(self.document, content_valid)
            tencent_util.add_quotation(self.document, quotation)
            tencent_util.add_picture(self.document, story)
            tencent_util.add_video(self.document, story)
            print('时间：', item[2])
            tencent_util.add_time(self.document, item[2])
            tencent_util.add_location(self.document, story)
            
                
 
    def start(self):
        while self.have_next_page:
            print('开始分析腾讯微博第　%s 页...' % self.page_index)
            page = weibo.get_next_page()
            print('是否有下一页：', self.have_next_page)
            print('下一页的地址： ', self.next_page_url)
            
            if self.page_index < tencent_weibo.START_PAGE_INDEX:
                print('>>>>>>>>>>>>>>>跳过本页<<<<<<<<<<<<<<<')
                self.page_index += 1
                continue
            if self.page_index >= tencent_weibo.END_PAGE_INDEX:
                self.have_next_page = False
            
            if self.page_index % tencent_weibo.SAVE_FILE_PAGE == 0:
                self.document.save('tencent_weibo_' 
                                   + str(self.page_index - tencent_weibo.SAVE_FILE_PAGE + 1)
                                   + '_' + str(self.page_index) + '.docx')
                self.document = Document()
            
            stories = self.get_stories(page)
            self.get_items(stories)
            
            self.page_index += 1
            time.sleep(1)
        else:
            self.page_index -= 1
            self.document.save('tencent_weibo_' + str(self.page_index // tencent_weibo.SAVE_FILE_PAGE + 1) + '_' + str(self.page_index) + '.docx')
            print('腾讯微博分析完成，共计　%s 页。' % self.page_index)
            print('备份了其中的第  %s 页到第　%s 页。' % (tencent_weibo.START_PAGE_INDEX,
                                             (tencent_weibo.END_PAGE_INDEX, self.page_index)[tencent_weibo.END_PAGE_INDEX > self.page_index]))

weibo = tencent_weibo('renminwangcom')
weibo.start()      