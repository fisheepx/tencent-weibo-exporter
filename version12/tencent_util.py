# -*- coding: utf-8 -*-
'''

Created on 2017/08/22

@author: yuyang
'''

import os
import urllib
import uuid
import re
import docx_ext

from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches

JPEG_EXTENSION = '.jpg'
PNG_EXTENSION = '.png'
GIF_EXTENSION = '.gif'
SPLIT_STRING = '///'
TOPIC_STRING = 'TTOOPPIICC'
EMOJI_STRING = 'EEMMOOJJII'
FRIEND_STRING = 'FFRRIIEENNDD'
URL_STRING = 'UURRLL'
QQEMO_STRING = 'QQEEMMOO'

def add_author(document, author):
    para = document.add_paragraph()
    run = para.add_run(author)
    font = run.font
    #font.name = 'Microsoft YaHei'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0x43, 0x6E, 0xEE)
    
def add_content(document, content, para = None, font_size = 16):
    if content.__contains__('k.t.qq.com'):
        pattern = re.compile(r'(<a href="http://k.t.qq.com.*?</a>)', re.S)
        topics = re.findall(pattern, content)
        for topic in topics:
            topic_word = topic.split('#')[1]
            content = content.replace(topic, SPLIT_STRING + TOPIC_STRING + '#' + topic_word + '#' + SPLIT_STRING)
            
    if content.__contains__('www/mb/images/emoji'):
        pattern_emoji = re.compile(r'(<img.*?>)', re.S)
        pattern_emoji_img = re.compile(r"crs='(.*?)'", re.S)
        emojis = re.findall(pattern_emoji, content)
        for emoji in emojis:
            emoji_url = re.findall(pattern_emoji_img, emoji)[0]
            filename = download_pic(emoji_url, PNG_EXTENSION)
            content = content.replace(emoji, SPLIT_STRING + EMOJI_STRING + filename + SPLIT_STRING)
    
    if content.__contains__('em rel="@'):
        pattern_friend = re.compile(r'(<em rel=.*?</em>)', re.S)
        pattern_friend_name = re.compile(r'<em.*?title="(.*?)"', re.S)
        friends = re.findall(pattern_friend, content)
        for friend in friends:
            friend_name = re.findall(pattern_friend_name, friend)[0]
            content = content.replace(friend, SPLIT_STRING + FRIEND_STRING + friend_name + SPLIT_STRING)
    
    if content.__contains__('http://url.cn'):
        pattern_url = re.compile(r'(<a href=.*?</a>)', re.S)
        pattern_url_str = re.compile(r'<a href="(.*?)"', re.S)
        urls = re.findall(pattern_url, content)
        for url in urls:
            url_str = re.findall(pattern_url_str, url)[0]
            content = content.replace(url, SPLIT_STRING + URL_STRING + url_str + SPLIT_STRING)
    
                
    if content.__contains__('www/mb/images/face'):
        pattern_qqemo = re.compile(r'(<img.*?>)', re.S)
        pattern_qqemo_img = re.compile(r"crs='(.*?)'", re.S)
        qqemos = re.findall(pattern_qqemo, content)
        for qqemo in qqemos:
            qqemo_url = re.findall(pattern_qqemo_img, qqemo)[0]
            filename = download_pic(qqemo_url, GIF_EXTENSION)
            content = content.replace(qqemo, SPLIT_STRING + QQEMO_STRING + filename + SPLIT_STRING)
        
    content_parts = content.split(SPLIT_STRING)
    
    if not para:
        para = document.add_paragraph()
    for content_part in content_parts:
        # delete first <div> mark 
        if content_part.startswith('<div>'):
            content_part = content_part[5:]
            
        if content_part.startswith(TOPIC_STRING):
            run = para.add_run(content_part.replace(TOPIC_STRING, ''))
            font = run.font
            font.italic = True
            font.bold = False
            font.size = Pt(font_size)
            font.color.rgb = RGBColor(0x00, 0x00, 0xCD)
        elif content_part.startswith(EMOJI_STRING):
            run = para.add_run()
            filename = content_part.replace(EMOJI_STRING, '')
            run.add_picture(filename)
        elif content_part.startswith(FRIEND_STRING):
            run = para.add_run(content_part.replace(FRIEND_STRING, ''))
            font = run.font
            font.italic = True
            font.bold = False
            font.size = Pt(font_size - 2)
            font.color.rgb = RGBColor(0xFF, 0x45, 0x00)
        elif content_part.startswith(URL_STRING):
            docx_ext.add_hyperlink(para, content_part.replace(URL_STRING, ''), 
                                   content_part.replace(URL_STRING, ''), '1E90FF', True)
        elif content_part.startswith(QQEMO_STRING):
            run = para.add_run()
            filename = content_part.replace(QQEMO_STRING, '')
            run.add_picture(filename)
        else:
            content_part = content_part.replace('&amp;', '&')
            content_part = content_part.replace('&gt;', '>')
            run = para.add_run(content_part)
            font = run.font
            font.bold = False
            font.size = Pt(font_size)
            font.color.rgb = RGBColor(0x08, 0x08, 0x08)

def add_quotation(document, quotation):
    if not quotation:
        return
    quotation_items = analyze_quotation(quotation)
    para = document.add_paragraph(style='IntenseQuote')
    if len(quotation_items) == 1:
        run = para.add_run(quotation_items[0])
        font = run.font
        font.bold = False
        font.size = Pt(12)
        font.color.rgb = RGBColor(0xA9, 0xA9, 0xA9)
        return
    run = para.add_run(quotation_items[0] + u'：')
    font = run.font
    font.bold = False
    font.size = Pt(12)
    font.color.rgb = RGBColor(0x48, 0xD1, 0xCC)
    add_content(document, quotation_items[1] + '\n', para, 12)
    
    filenames = analyze_pic(quotation)
    for filename in filenames:
        run_pic = para.add_run()
        run_pic.add_picture(filename, width=Inches(3))
        para.add_run('\n')
        
    run_time = para.add_run(quotation_items[2])
    font_time = run_time.font
    font_time.bold = False
    font_time.size = Pt(8)
    font_time.color.rgb = RGBColor(0x69, 0x69, 0x69)
    

def add_picture(document, story):
    filenames = analyze_pic(story)
    for filename in filenames:
        try:
            document.add_picture(filename, width=Inches(5))
        except:
            print '插入图片出错：' + filename

def add_time(document, time):
    para = document.add_paragraph()
    run = para.add_run(time)
    font = run.font
    font.italic = True
    #font.name = 'Microsoft YaHei'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x7A, 0x7A, 0x7A)
    
def add_location(document, story):
    location_items = analyze_loc(story)
    if len(location_items) <= 0:
        return
    link_name = location_items[2]
    google_map_url = 'https://maps.google.com/maps?q=' + location_items[0] + ',' + location_items[1]
    print google_map_url
    para = document.add_paragraph()
    run = para.add_run(u'位置：')
    font = run.font
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x7A, 0x7A, 0x7A)
    docx_ext.add_hyperlink(para, google_map_url, link_name, '4169E1', False)

def add_video(document, story):
    video_items = analyze_video(story)
    if not video_items:
        return
    para = document.add_paragraph()
    run = para.add_run()
    font = run.font
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x7A, 0x7A, 0x7A)
    docx_ext.add_hyperlink(para, video_items[0], video_items[1], '4169E1', False)
    document.add_picture(video_items[3], width=Inches(3))

def download_pic(url, extension):
    try:
        if not os.path.exists('.//pics'):
            os.mkdir('.//pics')
        filename = '.\\pics\\' + str(uuid.uuid4()) + extension
        urllib.urlretrieve(url, filename)
    except Exception:
        print 'Download picture error: ' + url
    return filename

def analyze_pic(story):
    filenames = []
    if story.__contains__('class="picBox"'):
        pattern = re.compile(r'<div class="picBox">\n<a href="(.*?)" data-like', re.S)
        img_url = re.findall(pattern, story)[0]
        print '图片：',
        print img_url
        filename = download_pic(img_url, JPEG_EXTENSION)
        filenames.append(filename)
    elif story.__contains__('class="tl_imgGroup'):
        pattern = re.compile(r'<div class="tl_imgGroup(.*?)<div class="miniMultiMedia clear"', re.S)
        imgs_str = re.findall(pattern, story)[0]
        pattern_img = re.compile(r'<a href="(.*?)" class="tl_imgGroup', re.S)
        imgs = re.findall(pattern_img, imgs_str)
        for img_url in imgs:
            print '图片：',
            print img_url
            filename = download_pic(img_url, JPEG_EXTENSION)
            filenames.append(filename)
            
    return filenames

def analyze_loc(story):
    location_items = []
    if story.__contains__('class="areaInfo"'):
        pattern = re.compile(r'boss="btn_check_tweetNear".*?lat=(.*?)&lng=(.*?)&addr=(.*?)" target', re.S)
        location_items = re.findall(pattern, story)[0]
        print u'位置：' + location_items[2]
        print u'经度：' + location_items[0]
        print u'纬度：' + location_items[1]

    return location_items

def analyze_video(story):
    video_items = []
    if story.__contains__('class="videoBox"'):
        pattern = re.compile(r'<div class="videoBox".*?realurl="(.*?)".*?reltitle="(.*?)".*?<img.*?crs="(.*?)"', re.S)
        video_items = re.findall(pattern, story)[0]
        print u'视频名称：' + video_items[1]
        print u'视频网址：' + video_items[0]
        print u'视频封面：' + video_items[2]
        try:
            filename = download_pic(video_items[2], '.jpg')
        except:
            u'下载视频封面出错：' + video_items[2]
            filename = None
        video_items = list(video_items)
        video_items.append(filename)
    return video_items

def depart_quotation(story):
    quotation_block = None
    if story.__contains__('class="replyBox"'):
        if story.__contains__('class="noMSource"'):#原文已被作者删除的情况
            pattern = re.compile(r'(<div class="replyBox".*?<div class="noMSource".*?</div>.*?</div>)', re.S)
            quotation_block = re.findall(pattern, story)[0]
            
        else:
            pattern = re.compile(r'(<div class="replyBox".*?<div class="msgBox".*?title=".*?" gender=' +
                         '.*?<div class="pubInfo.*?from="\\d*">.*?</a>.*?</div>.*?</div>)', re.S)
            quotation_block = re.findall(pattern, story)[0]
        story = story.replace(quotation_block, '')
    return story, quotation_block

def analyze_quotation(quotation):
    quotation_items = []
    if quotation.__contains__('class="noMSource"'):
        quotation_items = [u'原文已经被作者删除。']
        return quotation_items
    
    pattern = re.compile(r'<div class="msgCnt".*?title="(.*?)" gender=' +
                         '.*?<div>(.*?)</div>' +
                         '.*?<div class="pubInfo.*?from="\\d*">(.*?)</a>', re.S)
    quotation_items = re.findall(pattern, quotation)[0]
    print '原帖作者：', quotation_items[0]
    print '原帖内容：', quotation_items[1]
    print '原帖时间：', quotation_items[2]
    
    return quotation_items
    