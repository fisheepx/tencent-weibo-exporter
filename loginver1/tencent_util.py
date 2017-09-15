# -*- coding: utf-8 -*-
'''

Created on 2017/09/14

@author: yuyang
'''

from docx.shared import Pt
from docx.shared import RGBColor

JPEG_EXTENSION = '.jpg'
PNG_EXTENSION = '.png'
GIF_EXTENSION = '.gif'
SPLIT_STRING = '///'

def add_author(document, author):
    para = document.add_paragraph()
    run = para.add_run(author)
    font = run.font
    #font.name = 'Microsoft YaHei'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0x43, 0x6E, 0xEE)
    
def add_content(document, content, para = None, font_size = 16):
    if not para:
        para = document.add_paragraph()

    run = para.add_run(content)
    font = run.font
    font.bold = False
    font.size = Pt(font_size)
    font.color.rgb = RGBColor(0x08, 0x08, 0x08)    

def add_time(document, time):
    para = document.add_paragraph()
    run = para.add_run(time)
    font = run.font
    font.italic = True
    #font.name = 'Microsoft YaHei'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x7A, 0x7A, 0x7A)
