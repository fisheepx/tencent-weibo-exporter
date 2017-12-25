# -*- coding: utf-8 -*-
'''
change log:

    修改运行环境为Python3.x
    以对应控制台执行时中文LOG乱码导致异常
    =============OLD=============
    对应含有引用的内容
        在非login版本基础上增加了含有视频的内容
    修改点击下一页的逻辑
    使用非login版代码对应含有主题，好友，Emoji，链接等全部内容
    添加位置信息
    对应带有视频有内容
    添加添加保存图片功能
    保存纯文本内容
    

Created on 2017/12/22

@author: yuyang
'''
from time import sleep
from selenium import webdriver
from docx import Document
import tencent_util

class tencent_weibo:
    '''
    Tencent weibo object.
    '''
    START_PAGE_INDEX = 1
    END_PAGE_INDEX = 100
    SAVE_FILE_PAGE = 10


    def __init__(self):
        '''
        Constructor
        '''
        self.page_index = 1
        self.have_next_page = True
        self.next_page_url = None
        self.stories = None
        self.browser = webdriver.Firefox()
        self.document = Document()
        
    def login(self):
        print('开始登录腾讯微博...')
        url = 'http://t.qq.com/?lang=zh_CN'
        self.browser.get(url)
        sleep(3)
        try:
            self.browser.switch_to.frame('login_div')
            self.browser.find_element_by_class_name('face').click()

            print('腾讯微博登录完毕...')
        except:
            print('登录超时...')

        self.browser.set_page_load_timeout(5)
        sleep(10)

        try:
            current_url = self.browser.current_url
            #防止找不到对象异常
            self.browser.switch_to.default_content()
            self.browser.get(current_url + '/mine')
            sleep(10)
            print('打开我的广播...')
        except:
            print('打开我的广播超时...')
    
    def get_stories(self):
        talk_list = self.browser.find_element_by_id('talkList')
        self.stories = talk_list.find_elements_by_tag_name('li') 

    def get_items(self):
        for story in self.stories:
            story, quotation = tencent_util.depart_quotation(self.browser, story)
            
            author = story.find_element_by_class_name('userName').find_element_by_tag_name('a').get_attribute('title')
            print('作者：', author)
            content_html = story.find_element_by_class_name('msgCnt').get_attribute('innerHTML')
            content_valid = ''.join(c for c in content_html if ord(c) >= 32)# delete control chars
            content = story.find_element_by_class_name('msgCnt').text
            print('内容：', content)
            time = story.find_element_by_class_name('time').text
            print('时间：', time)
            print('----------------------------------------------------------------------------------')
            self.document.add_heading('', 0)
            tencent_util.add_author(self.document, author)
            tencent_util.add_content(self.document, content_valid)
            tencent_util.add_quotation(self.document, quotation)
            tencent_util.add_picture(self.document, story)
            tencent_util.add_video(self.document, story)
            tencent_util.add_time(self.document,time)
            tencent_util.add_location(self.document, story)
 
    def click_next_page(self):
        pageNav = self.browser.find_element_by_id('pageNav')
        try:
            next_page = pageNav.find_element_by_link_text(u'下一页')
        except:
            self.have_next_page = False
            print('取得下一页链接超时...')
        try:
            self.next_page_url = next_page.get_attribute('href')
            print('是否有下一页：', self.have_next_page)
            print('下一页的地址：',self.next_page_url)
            if self.next_page_url:
                next_page.click()
            else:
                None# never come to this line because it always timeout
        except:
            print('点击下一页超时...')
 
    def start(self):
        self.login()
        while self.have_next_page:
            print('开始分析腾讯微博第　%s 页...' % self.page_index)
            
            if self.page_index < tencent_weibo.START_PAGE_INDEX:
                print('>>>>>>>>>>>>>>>跳过本页<<<<<<<<<<<<<<<')
                self.click_next_page()
                self.page_index += 1
                continue
            if self.page_index >= tencent_weibo.END_PAGE_INDEX:
                self.have_next_page = False
            
            if self.page_index % tencent_weibo.SAVE_FILE_PAGE == 0:
                self.document.save('tencent_weibo_' 
                                   + str(self.page_index - tencent_weibo.SAVE_FILE_PAGE + 1)
                                   + '_' + str(self.page_index) + '.docx')
                self.document = Document()
            
            self.get_stories()
            self.get_items()
            self.click_next_page()
            
            self.page_index += 1
            sleep(1)
        else:
            self.page_index -= 1
            self.document.save('tencent_weibo_' 
                   + str(self.page_index // tencent_weibo.SAVE_FILE_PAGE + 1)
                   + '_' + str(self.page_index) + '.docx')
            print('腾讯微博分析完成，共计　%s 页。' % self.page_index)
            print('备份了其中的第  %s 页到第　%s 页。' % (tencent_weibo.START_PAGE_INDEX,
                                             (tencent_weibo.END_PAGE_INDEX, self.page_index)[tencent_weibo.END_PAGE_INDEX > self.page_index]))
        
        self.browser.quit()
        
weibo = tencent_weibo()
weibo.start()      