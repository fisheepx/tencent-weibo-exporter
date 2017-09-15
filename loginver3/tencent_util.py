# -*- coding: utf-8 -*-
'''

Created on 2017/09/14

@author: yuyang
'''

import os
import urllib
import uuid
import re
import docx_ext

from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches

JPEG_EXTENSION = '.jpg'
PNG_EXTENSION = '.png'
GIF_EXTENSION = '.gif'
SPLIT_STRING = '///'

def add_author(document, author):
    para = document.add_paragraph()
    run = para.add_run(author)
    font = run.font
    #font.name = 'Microsoft YaHei'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0x43, 0x6E, 0xEE)
    
def add_content(document, content, para = None, font_size = 16):
    if not para:
        para = document.add_paragraph()

    run = para.add_run(content)
    font = run.font
    font.bold = False
    font.size = Pt(font_size)
    font.color.rgb = RGBColor(0x08, 0x08, 0x08)    

def add_picture(document, story):
    filenames = analyze_pic(story)
    for filename in filenames:
        try:
            document.add_picture(filename, width=Inches(5))
        except:
            print '插入图片出错：' + filename

def add_time(document, time):
    para = document.add_paragraph()
    run = para.add_run(time)
    font = run.font
    font.italic = True
    #font.name = 'Microsoft YaHei'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x7A, 0x7A, 0x7A)
    
def add_location(document, story):
    location_items = analyze_loc(story)
    if len(location_items) <= 0:
        return
    link_name = location_items[2]
    google_map_url = 'https://maps.google.com/maps?q=' + location_items[0] + ',' + location_items[1]
    para = document.add_paragraph()
    run = para.add_run(u'位置：')
    font = run.font
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x7A, 0x7A, 0x7A)
    docx_ext.add_hyperlink(para, google_map_url, link_name, '4169E1', False)

def add_video(document, story):
    video_items = analyze_video(story)
    if not video_items:
        return
    para = document.add_paragraph()
    run = para.add_run()
    font = run.font
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x7A, 0x7A, 0x7A)
    docx_ext.add_hyperlink(para, video_items[0], video_items[1], '4169E1', False)
    try:
        document.add_picture(video_items[3], width=Inches(3))
    except:
        print '视频封面插入出错：' + video_items[3]
   

def download_pic(url, extension):
    try:
        if not os.path.exists('.//pics'):
            os.mkdir('.//pics')
        filename = '.\\pics\\' + str(uuid.uuid4()) + extension
        urllib.urlretrieve(url, filename)
    except Exception:
        print '下载图片出错： ' + url
    return filename

def analyze_pic(story):
    filenames = []
    picBox = None
    imgGroup = None
    try:
        picBox = story.find_element_by_class_name('picBox')
    except:
        None
    try:
        imgGroup = story.find_element_by_class_name('tl_imgGroup')
    except:
        None
    
    if picBox:# one picture
        img_url = picBox.find_element_by_tag_name('a').get_attribute('href')
        print '图片：', img_url
        filename = download_pic(img_url, JPEG_EXTENSION)
        filenames.append(filename)

    elif imgGroup:# multi picture
        a_tags = imgGroup.find_elements_by_tag_name('a')
        for a_tag in a_tags:
            img_url = a_tag.get_attribute('href')
            print '图片：', img_url
            filename = download_pic(img_url, JPEG_EXTENSION)
            filenames.append(filename)
            
    return filenames

def analyze_loc(story):
    location_items = []
    areaInfo = None
    try:
        areaInfo = story.find_element_by_class_name('areaInfo')
    except:
        None
        
    if areaInfo:
        pattern = re.compile(r'http.*?lat=(.*?)&amp;lng=(.*?)&amp;addr=(.*?)"', re.S)
        location_items = re.findall(pattern, areaInfo.get_attribute('innerHTML'))[0]
        print '位置：', location_items[2]
        print '经度：', location_items[0]
        print '纬度：', location_items[1]

    return location_items

def analyze_video(story):
    video_items = []
    videoBox = None
    try:
        videoBox = story.find_element_by_class_name('videoBox')
    except:
        None
        
    if videoBox:
        pattern = re.compile(r'realurl="(.*?)".*?reltitle="(.*?)".*?<img.*?crs="(.*?)"', re.S)
        video_items = re.findall(pattern, videoBox.get_attribute('outerHTML'))[0]
        print '视频名称：', video_items[1]
        print '视频网址：', video_items[0]
        print '视频封面：', video_items[2]
        try:
            filename = download_pic(video_items[2], JPEG_EXTENSION)
        except:
            print '下载视频封面出错：', video_items[2]
            filename = None
        video_items = list(video_items)
        video_items.append(filename)
    return video_items