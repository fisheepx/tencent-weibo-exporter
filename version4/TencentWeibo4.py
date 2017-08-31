# -*- coding: utf-8 -*-
'''
change log:

    优化生成Word文件格式
    下载图片时文件夹不存在则创建
    =============OLD=============
    对应一条微博多张图片时的图片下载
    对应含有图片的微博
    备份微博内容（仅文字部分）  

Created on 2017/08/10

@author: yuyang
'''
import time
import urllib2
import re
from docx import Document
import TencentUtil

class TencentWeibo:
    '''
    Tencent weibo object.
    '''

    def __init__(self, id):
        '''
        Constructor
        '''
        self.id = id
        self.pageIndex = 1
        self.haveNextPage = True
        self.nextPageUrl = None
        self.userAgent = 'Mozilla/4.0 (compatible; MSIE 5.5; Windows NT)'
        self.headers = {'User-Agent' : self.userAgent}
        self.stories = []
        self.document = Document()
        
    def getNextPage(self):
        if self.pageIndex == 1:
            url = 'http://t.qq.com/' + self.id + '?mode=0&lang=zh_CN'
        else:
            url = self.nextPageUrl + '&lang=zh_CN'
        try:
            request = urllib2.Request(url, headers = self.headers)
            response = urllib2.urlopen(request)
            return response.read().decode('utf-8')
        except urllib2.URLError, e:
            if hasattr(e, 'reason'):
                print '腾讯微博第 %d 页连接失败： %s' % (self.pageIndex, e.reason)
                return None
 
    def getStories(self, page):
        pattern = re.compile(r'<ul id="talkList"(.*?)</ul>', re.S)
        
        stories_str = re.findall(pattern, page)
        stories = stories_str[0].split('<li')
        
        for match in re.finditer(r'<a href="\?mode=0&.*?"', page):
            pass
        self.nextPageUrl = match.group().replace('<a href="', 'http://t.qq.com/' + self.id)[:-1]
        
        if not page.__contains__(u'下一页'):
            self.haveNextPage = False
            
        return stories
            
    def getItems(self, stories):
        
        pattern = re.compile(r'<div class="msgBox".*?<div class="userName".*?title="(.*?)" gender=' +
                             '.*?<div class="msgCnt">(.*?)</div>' +
                             '.*?<div class="pubInfo.*?from="\\d*">(.*?)</a>', re.S)
        
        for story in stories:
            items = re.findall(pattern, story)
            if not len(items):
                continue
            
            item = items[0]
            print '作者：', item[0]
            self.document.add_heading('', 0)
            TencentUtil.addAuthor(self.document, item[0])
            print '内容：', item[1]
            content_valid = ''.join(c for c in item[1] if ord(c) >= 32)# delete control chars
            TencentUtil.addContent(self.document, content_valid)
            filenames = TencentUtil.analyzePic(story)
            for filename in filenames:
                try:
                    self.document.add_picture(filename)
                except:
                    print '插入图片出错：' + filename
            print '时间：', item[2]
            TencentUtil.addTime(self.document, item[2])
            
                
 
    def start(self):
        while self.haveNextPage:
            print '开始分析腾讯微博第　%s 页...' % self.pageIndex
            page = weibo.getNextPage()
            #print page
            stories = self.getStories(page)
            self.getItems(stories)
            print '是否有下一页：', self.haveNextPage
            print '下一页的地址： ', self.nextPageUrl
            self.pageIndex += 1
            #if self.pageIndex > 1:#test code
            #    self.document.save('tencent.docx')#test code
            #    break#test code
            time.sleep(1)
        else:
            print '腾讯微博备份完成，共计　%s 页' % self.pageIndex
            self.document.save('tencent.docx')
        
weibo = TencentWeibo('renminwangcom')
weibo.start()      