# -*- coding: utf-8 -*-
'''
change log:

    备份微博内容（仅文字部分）
    =============OLD=============
    
Created on 2017/08/10

@author: yuyang
'''
import urllib2
import re
from docx import Document

class TencentWeibo:
    '''
    Tencent weibo object.
    '''

    def __init__(self, id):
        '''
        Constructor
        '''
        self.id = id
        self.pageIndex = 1
        self.haveNextPage = True
        self.nextPageUrl = None
        self.userAgent = 'Mozilla/4.0 (compatible; MSIE 5.5; Windows NT)'
        self.headers = {'User-Agent' : self.userAgent}
        self.stories = []
        self.document = Document()
        
    def getNextPage(self):
        if self.pageIndex == 1:
            url = 'http://t.qq.com/' + self.id + '?mode=0&lang=zh_CN'
        else:
            url = self.nextPageUrl + '&lang=zh_CN'
        try:
            request = urllib2.Request(url, headers = self.headers)
            response = urllib2.urlopen(request)
            return response.read().decode('utf-8')
        except urllib2.URLError, e:
            if hasattr(e, 'reason'):
                print '腾讯微博第 %d 页连接失败： %s' % (self.pageIndex, e.reason)
                return None
 
    def getItems(self, page):
        #pattern = re.compile('<div class="msgBox".*?target="_blank">(.*?)</a>', re.S)
        #pattern = re.compile(r'<div class="msgBox".*?target="_blank">(.*?)</a>' +
        #                     '.*?<div class="msgBox".*?<div class="msgCnt">(.*?)</div>' +
        #                     '.*?<div class="pubInfo.*?from="\\d*">(.*?)</a>', re.S)
        pattern = re.compile(r'<div class="userName".*?title="(.*?)" gender=' +
                             '.*?<div class="msgCnt">(.*?)</div>' +
                             '.*?<div class="pubInfo.*?from="\\d*">(.*?)</a>', re.S)
        items = re.findall(pattern, page)
        for item in items:
            print '作者：', item[0]
            self.document.add_heading('', 0)
            self.document.add_heading(item[0], level=1)
            print '内容：', item[1]
            content_valid = ''.join(c for c in item[1] if ord(c) >= 32)# delete control chars
            self.document.add_paragraph(content_valid, style='ListBullet')
            print '时间：', item[2]
            p = self.document.add_paragraph('', style='ListBullet')
            p.add_run(item[2]).italic = True
            #self.document.add_page_break()
 
        for match in re.finditer(r'<a href="\?mode=0&.*?"', page):
            pass
        self.nextPageUrl = match.group().replace('<a href="', 'http://t.qq.com/' + self.id)[:-1]
        #print page
        try:
            page.index(u'下一页')
        except ValueError:
            self.haveNextPage = False
            

 
    def start(self):
        while self.haveNextPage:
            print '开始分析腾讯微博第　%s 页...' % self.pageIndex
            page = weibo.getNextPage()
            #print page
            self.getItems(page)
            print '是否有下一页：', self.haveNextPage
            print '得到下一页地址： ', self.nextPageUrl
            self.pageIndex += 1
            #if self.pageIndex > 4:#test code
                #self.document.save('tencent.docx')#test code
                #break#test code
        else:
            print '腾讯微博备份完成，共计　%s 页' % self.pageIndex
            self.document.save('tencent.docx')
        
weibo = TencentWeibo('renminwangcom')
weibo.start()      