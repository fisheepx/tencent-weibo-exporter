# -*- coding: utf-8 -*-
'''
Created on 2017/08/22

@author: yuyang
'''

import os
import urllib
import uuid
import re

from docx.shared import Pt
from docx.shared import RGBColor

def downloadPic(url):
    try:
        if not os.path.exists('.//pics'):
            os.mkdir('.//pics')
        filename = '.\\pics\\' + str(uuid.uuid4()) + '.jpg'
        urllib.urlretrieve(url, filename)
    except Exception:
        print 'Download picture error: ' + url
    return filename

def analyzePic(story):
    filenames = []
    if story.__contains__('class="picBox"'):
        pattern = re.compile(r'<div class="picBox">\n<a href="(.*?)" data-like', re.S)
        img_url = re.findall(pattern, story)[0]
        print '图片：',
        print img_url
        filename = downloadPic(img_url)
        filenames.append(filename)
    elif story.__contains__('class="tl_imgGroup'):
        pattern = re.compile(r'<div class="tl_imgGroup(.*?)<div class="miniMultiMedia clear"', re.S)
        imgs_str = re.findall(pattern, story)[0]
        pattern_img = re.compile(r'<a href="(.*?)" class="tl_imgGroup', re.S)
        imgs = re.findall(pattern_img, imgs_str)
        for img_url in imgs:
            print '图片：',
            print img_url
            filename = downloadPic(img_url)
            filenames.append(filename)
            
    return filenames

def addAuthor(document, author):
    para = document.add_paragraph()
    run = para.add_run(author)
    font = run.font
    #font.name = 'Microsoft YaHei'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0x43, 0x6E, 0xEE)
    
def addContent(document, content):
    para = document.add_paragraph()
    run = para.add_run(content)
    font = run.font
    #font.name = 'Microsoft YaHei'
    font.size = Pt(16)
    font.color.rgb = RGBColor(0x08, 0x08, 0x08)
    
def addTime(document, time):
    para = document.add_paragraph()
    run = para.add_run(time)
    font = run.font
    font.italic = True
    #font.name = 'Microsoft YaHei'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x7A, 0x7A, 0x7A)
        
    