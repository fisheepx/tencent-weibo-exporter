# -*- coding: utf-8 -*-
'''
change log:

    添加添加保存图片功能
    =============OLD=============
    保存纯文本内容
    

Created on 2017/09/14

@author: yuyang
'''
from time import sleep
from selenium import webdriver
from docx import Document
import tencent_util

class tencent_weibo:
    '''
    Tencent weibo object.
    '''
    START_PAGE_INDEX = 1
    END_PAGE_INDEX = 1
    SAVE_FILE_PAGE = 20


    def __init__(self):
        '''
        Constructor
        '''
        self.page_index = 1
        self.have_next_page = True
        self.next_page_url = None
        self.stories = None
        self.browser = webdriver.Firefox()
        self.document = Document()
        
    def login(self):
        url = 'http://t.qq.com/?lang=zh_CN'
        self.browser.get(url)
        try:
            self.browser.switch_to_frame('login_div')
            self.browser.find_element_by_class_name('face').click()
            print '腾讯微博登录完毕...'
        except:
            print '登录超时...'
        sleep(3)
    
    def next_page(self):
        if self.page_index == 1:
            self.browser.set_page_load_timeout(5)
            try:
                current_url = self.browser.current_url
                self.browser.get(current_url + '/mine')
                print '打开我的广播...'
            except:
                print '打开我的广播超时...'
        else:
            try:
                pageNav = self.browser.find_element_by_id('pageNav')
                next_page = pageNav.find_element_by_link_text(u'下一页')
                if not next_page:
                    self.have_next_page = False
                    return
                print '下一页地址：',
                print next_page.get_attribute('href')
                next_page.click()
            except:
                print '点击下一页超时...'
                
        talk_list = self.browser.find_element_by_id('talkList')
        self.stories = talk_list.find_elements_by_tag_name('li') 

    def get_items(self):
        for story in self.stories:
            author = story.find_element_by_class_name('userName').find_element_by_tag_name('a').get_attribute('title')
            #print story.find_element_by_class_name('userName').get_attribute('innerHTML')
            print '作者：', author
            content = story.find_element_by_class_name('msgCnt').text
            print '内容：', content
            time = story.find_element_by_class_name('time').text
            print '时间：', time
            print '----------------------------------------------------------------------------------'
            self.document.add_heading('', 0)
            tencent_util.add_author(self.document, author)
            tencent_util.add_content(self.document, content)
            tencent_util.add_picture(self.document, story)
            tencent_util.add_time(self.document,time)
 
    def start(self):
        self.login()
        while self.have_next_page:
            print '开始分析腾讯微博第　%s 页...' % self.page_index
            self.next_page()
            print '是否有下一页：', self.have_next_page
            print '下一页的地址： ', self.next_page_url
            
            if self.page_index < tencent_weibo.START_PAGE_INDEX:
                print '>>>>>>>>>>>>>>>跳过本页<<<<<<<<<<<<<<<'
                self.page_index += 1
                continue
            if self.page_index >= tencent_weibo.END_PAGE_INDEX:
                self.have_next_page = False
            
            if self.page_index % tencent_weibo.SAVE_FILE_PAGE == 0:
                self.document.save('tencent_weibo_' 
                                   + str(self.page_index - tencent_weibo.SAVE_FILE_PAGE + 1)
                                   + '_' + str(self.page_index) + '.docx')
                self.document = Document()
            
            self.get_items()
            
            self.page_index += 1
            sleep(1)
        else:
            self.page_index -= 1
            self.document.save('tencent_weibo_' 
                   + str(self.page_index / tencent_weibo.SAVE_FILE_PAGE * tencent_weibo.SAVE_FILE_PAGE + 1)
                   + '_' + str(self.page_index) + '.docx')
            print '腾讯微博分析完成，共计　%s 页。' % self.page_index
            print '备份了其中的第  %s 页到第　%s 页。' % (tencent_weibo.START_PAGE_INDEX, 
                                             (tencent_weibo.END_PAGE_INDEX, self.page_index)[tencent_weibo.END_PAGE_INDEX > self.page_index])
        
        self.browser.quit()
        
weibo = tencent_weibo()
weibo.start()      