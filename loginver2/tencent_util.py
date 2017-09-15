# -*- coding: utf-8 -*-
'''

Created on 2017/09/14

@author: yuyang
'''

import os
import urllib
import uuid

from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches

JPEG_EXTENSION = '.jpg'
PNG_EXTENSION = '.png'
GIF_EXTENSION = '.gif'
SPLIT_STRING = '///'

def add_author(document, author):
    para = document.add_paragraph()
    run = para.add_run(author)
    font = run.font
    #font.name = 'Microsoft YaHei'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0x43, 0x6E, 0xEE)
    
def add_content(document, content, para = None, font_size = 16):
    if not para:
        para = document.add_paragraph()

    run = para.add_run(content)
    font = run.font
    font.bold = False
    font.size = Pt(font_size)
    font.color.rgb = RGBColor(0x08, 0x08, 0x08)    

def add_picture(document, story):
    filenames = analyze_pic(story)
    for filename in filenames:
        try:
            document.add_picture(filename, width=Inches(5))
        except:
            print '插入图片出错：' + filename

def add_time(document, time):
    para = document.add_paragraph()
    run = para.add_run(time)
    font = run.font
    font.italic = True
    #font.name = 'Microsoft YaHei'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x7A, 0x7A, 0x7A)

def download_pic(url, extension):
    try:
        if not os.path.exists('.//pics'):
            os.mkdir('.//pics')
        filename = '.\\pics\\' + str(uuid.uuid4()) + extension
        urllib.urlretrieve(url, filename)
    except Exception:
        print '下载图片出错： ' + url
    return filename

def analyze_pic(story):
    filenames = []
    picBox = None
    imgGroup = None
    try:
        picBox = story.find_element_by_class_name('picBox')
    except:
        None
    try:
        imgGroup = story.find_element_by_class_name('tl_imgGroup')
    except:
        None
    
    if picBox:# one picture
        img_url = picBox.find_element_by_tag_name('a').get_attribute('href')
        print '图片：', img_url
        filename = download_pic(img_url, JPEG_EXTENSION)
        filenames.append(filename)

    elif imgGroup:# multi picture
        a_tags = imgGroup.find_elements_by_tag_name('a')
        for a_tag in a_tags:
            img_url = a_tag.get_attribute('href')
            print '图片：', img_url
            filename = download_pic(img_url, JPEG_EXTENSION)
            filenames.append(filename)
            
    return filenames