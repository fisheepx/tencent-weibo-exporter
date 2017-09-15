# -*- coding: utf-8 -*-
'''

Created on 2017/09/14

@author: yuyang
'''

import os
import urllib
import uuid
import re
import docx_ext

from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches

JPEG_EXTENSION = '.jpg'
PNG_EXTENSION = '.png'
GIF_EXTENSION = '.gif'
SPLIT_STRING = '///'
TOPIC_STRING = 'TTOOPPIICC'
EMOJI_STRING = 'EEMMOOJJII'
FRIEND_STRING = 'FFRRIIEENNDD'
URL_STRING = 'UURRLL'
QQEMO_STRING = 'QQEEMMOO'
OTHEREMO_STRING = 'OOTTHHEERR'

def add_author(document, author):
    para = document.add_paragraph()
    run = para.add_run(author)
    font = run.font
    #font.name = 'Microsoft YaHei'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0x43, 0x6E, 0xEE)
    
def add_content(document, content, para = None, font_size = 16):
    print content
    if content.__contains__('k.t.qq.com'):
        pattern = re.compile(r'(<a href="http://k.t.qq.com.*?</a>)', re.S)
        topics = re.findall(pattern, content)
        for topic in topics:
            topic_word = topic.split('#')[1]
            content = content.replace(topic, SPLIT_STRING + TOPIC_STRING + '#' + topic_word + '#' + SPLIT_STRING)
            
    if content.__contains__('www/mb/images/emoji'):
        pattern_emoji = re.compile(r'(<img.*?>)', re.S)
        pattern_emoji_img = re.compile(r'crs="(.*?)"', re.S)
        emojis = re.findall(pattern_emoji, content)
        for emoji in emojis:
            emoji_url = re.findall(pattern_emoji_img, emoji)[0]
            filename = download_pic(emoji_url, PNG_EXTENSION)
            content = content.replace(emoji, SPLIT_STRING + EMOJI_STRING + filename + SPLIT_STRING)
    
    if content.__contains__('em rel="@'):
        pattern_friend = re.compile(r'(<em rel=.*?</em>)', re.S)
        pattern_friend_name = re.compile(r'<em.*?title="(.*?)"', re.S)
        friends = re.findall(pattern_friend, content)
        for friend in friends:
            friend_name = re.findall(pattern_friend_name, friend)[0]
            content = content.replace(friend, SPLIT_STRING + FRIEND_STRING + friend_name + SPLIT_STRING)
    
    if content.__contains__('http://url.cn'):
        pattern_url = re.compile(r'(<a href=.*?</a>)', re.S)
        pattern_url_str = re.compile(r'<a href="(.*?)"', re.S)
        urls = re.findall(pattern_url, content)
        for url in urls:
            url_str = re.findall(pattern_url_str, url)[0]
            content = content.replace(url, SPLIT_STRING + URL_STRING + url_str + SPLIT_STRING)
    
                
    if content.__contains__('www/mb/images/face'):
        pattern_qqemo = re.compile(r'(<img.*?>)', re.S)
        pattern_qqemo_img = re.compile(r'crs="(.*?)"', re.S)
        qqemos = re.findall(pattern_qqemo, content)
        for qqemo in qqemos:
            qqemo_url = re.findall(pattern_qqemo_img, qqemo)[0]
            filename = download_pic(qqemo_url, GIF_EXTENSION)
            content = content.replace(qqemo, SPLIT_STRING + QQEMO_STRING + filename + SPLIT_STRING)
    
    if content.__contains__('<img class='):
        pattern_other_emo = re.compile(r'(<img.*?>)', re.S)
        pattern_other_emo_img = re.compile(r'<img.*?crs=(.*?) title=', re.S)
        pattern_other_emo_img_only = re.compile(r'<img.*?crs=(.*?)>', re.S)
        pattern_other_emos = re.findall(pattern_other_emo, content)
        for other_emo in pattern_other_emos:
            other_emo_match = re.findall(pattern_other_emo_img, other_emo)
            if not other_emo_match:# some emoji have special pattern
                other_emo_match = re.findall(pattern_other_emo_img_only, other_emo)
            other_emo_url = other_emo_match[0]
            other_emo_url = other_emo_url[1:-1]# delete start and end mark ' "
            filename = download_pic(other_emo_url, other_emo_url[-4:])
            content = content.replace(other_emo, SPLIT_STRING + OTHEREMO_STRING + filename + SPLIT_STRING)
        
    content_parts = content.split(SPLIT_STRING)
    
    if not para:
        para = document.add_paragraph()
    for content_part in content_parts:
        # delete first <div> mark 
        if content_part.startswith('<div>'):
            content_part = content_part[5:]
        if content_part.endswith('</div>'):
            content_part = content_part[:-6]
            
        if content_part.startswith(TOPIC_STRING):
            run = para.add_run(content_part.replace(TOPIC_STRING, ''))
            font = run.font
            font.italic = True
            font.bold = False
            font.size = Pt(font_size)
            font.color.rgb = RGBColor(0x00, 0x00, 0xCD)
        elif content_part.startswith(EMOJI_STRING):
            run = para.add_run()
            filename = content_part.replace(EMOJI_STRING, '')
            run.add_picture(filename)
        elif content_part.startswith(FRIEND_STRING):
            run = para.add_run(content_part.replace(FRIEND_STRING, ''))
            font = run.font
            font.italic = True
            font.bold = False
            font.size = Pt(font_size - 2)
            font.color.rgb = RGBColor(0xFF, 0x45, 0x00)
        elif content_part.startswith(URL_STRING):
            docx_ext.add_hyperlink(para, content_part.replace(URL_STRING, ''), 
                                   content_part.replace(URL_STRING, ''), '1E90FF', True)
        elif content_part.startswith(QQEMO_STRING):
            run = para.add_run()
            filename = content_part.replace(QQEMO_STRING, '')
            run.add_picture(filename)
        elif content_part.startswith(OTHEREMO_STRING):
            run = para.add_run()
            filename = content_part.replace(OTHEREMO_STRING, '')
            run.add_picture(filename)
        else:
            content_part = content_part.replace('&amp;', '&')
            content_part = content_part.replace('&gt;', '>')
            content_part = content_part.replace('&quot;', '"')
            content_part = content_part.replace('&lt;', '<')
            run = para.add_run(content_part)
            font = run.font
            font.bold = False
            font.size = Pt(font_size)
            font.color.rgb = RGBColor(0x08, 0x08, 0x08)
  

def add_picture(document, story):
    filenames = analyze_pic(story)
    for filename in filenames:
        try:
            document.add_picture(filename, width=Inches(5))
        except:
            print '插入图片出错：' + filename

def add_time(document, time):
    para = document.add_paragraph()
    run = para.add_run(time)
    font = run.font
    font.italic = True
    #font.name = 'Microsoft YaHei'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x7A, 0x7A, 0x7A)
    
def add_location(document, story):
    location_items = analyze_loc(story)
    if len(location_items) <= 0:
        return
    link_name = location_items[2]
    google_map_url = 'https://maps.google.com/maps?q=' + location_items[0] + ',' + location_items[1]
    para = document.add_paragraph()
    run = para.add_run(u'位置：')
    font = run.font
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x7A, 0x7A, 0x7A)
    docx_ext.add_hyperlink(para, google_map_url, link_name, '4169E1', False)

def add_video(document, story):
    video_items = analyze_video(story)
    if not video_items:
        return
    para = document.add_paragraph()
    run = para.add_run()
    font = run.font
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x7A, 0x7A, 0x7A)
    docx_ext.add_hyperlink(para, video_items[0], video_items[1], '4169E1', False)
    try:
        document.add_picture(video_items[3], width=Inches(3))
    except:
        print '视频封面插入出错：' + video_items[3]
   

def download_pic(url, extension):
    try:
        if not os.path.exists('.//pics'):
            os.mkdir('.//pics')
        filename = '.\\pics\\' + str(uuid.uuid4()) + extension
        urllib.urlretrieve(url, filename)
    except Exception:
        print '下载图片出错： ' + url
    return filename

def analyze_pic(story):
    filenames = []
    picBox = None
    imgGroup = None
    try:
        picBox = story.find_element_by_class_name('picBox')
    except:
        None
    try:
        imgGroup = story.find_element_by_class_name('tl_imgGroup')
    except:
        None
    
    if picBox:# one picture
        img_url = picBox.find_element_by_tag_name('a').get_attribute('href')
        print '图片：', img_url
        filename = download_pic(img_url, JPEG_EXTENSION)
        filenames.append(filename)

    elif imgGroup:# multi picture
        a_tags = imgGroup.find_elements_by_tag_name('a')
        for a_tag in a_tags:
            img_url = a_tag.get_attribute('href')
            print '图片：', img_url
            filename = download_pic(img_url, JPEG_EXTENSION)
            filenames.append(filename)
            
    return filenames

def analyze_loc(story):
    location_items = []
    areaInfo = None
    try:
        areaInfo = story.find_element_by_class_name('areaInfo')
    except:
        None
        
    if areaInfo:
        pattern = re.compile(r'http.*?lat=(.*?)&amp;lng=(.*?)&amp;addr=(.*?)"', re.S)
        location_items = re.findall(pattern, areaInfo.get_attribute('innerHTML'))[0]
        print '位置：', location_items[2]
        print '经度：', location_items[0]
        print '纬度：', location_items[1]

    return location_items

def analyze_video(story):
    video_items = []
    videoBox = None
    try:
        videoBox = story.find_element_by_class_name('videoBox')
    except:
        None
        
    if videoBox:
        pattern = re.compile(r'realurl="(.*?)".*?reltitle="(.*?)".*?<img.*?crs="(.*?)"', re.S)
        video_items = re.findall(pattern, videoBox.get_attribute('outerHTML'))[0]
        print '视频名称：', video_items[1]
        print '视频网址：', video_items[0]
        print '视频封面：', video_items[2]
        try:
            filename = download_pic(video_items[2], JPEG_EXTENSION)
        except:
            print '下载视频封面出错：', video_items[2]
            filename = None
        video_items = list(video_items)
        video_items.append(filename)
    return video_items